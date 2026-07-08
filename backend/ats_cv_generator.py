"""
ATS CV Generator
================
Genera CVs en formato Word (.docx), PDF y texto plano optimizados para sistemas ATS (Applicant Tracking Systems).

Principios ATS:
- Sin tablas complejas ni columnas múltiples
- Fuentes estándar (Calibri/Arial)
- Sin encabezados/pies de página con info crítica
- Sin imágenes, íconos ni gráficos
- Texto plano y jerárquico con encabezados claros
- Sin cuadros de texto flotantes
- Márgenes estándar (1 pulgada)
"""

import re
import io
from typing import Optional, Dict, Any

# python-docx es opcional; si no está instalado, sólo devolvemos texto plano
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Try to import PDF generation libraries - use xhtml2pdf which is easier on Windows
try:
    from xhtml2pdf import pisa
    XHTML2PDF_AVAILABLE = True
except ImportError:
    XHTML2PDF_AVAILABLE = False


# ─── Sección Patterns ────────────────────────────────────────────────────────

_SECTION_KEYWORDS = {
    "experiencia": ["experiencia", "experience", "trabajo", "empleo", "laboral", "profesional"],
    "educacion":   ["educaci", "education", "académic", "formaci", "estudios", "universidad", "escuela"],
    "habilidades": ["habilidad", "skill", "competencia", "conocimiento", "tecnolog"],
    "certificaciones": ["certificaci", "certif", "curso", "diploma", "licencia"],
    "idiomas":     ["idioma", "language", "lenguaje"],
    "proyectos":   ["proyecto", "project"],
    "resumen":     ["resumen", "perfil", "summary", "objetivo", "profile", "acerca"],
    "contacto":    ["contacto", "contact", "información personal", "datos"],
}


def _detect_section(line: str) -> Optional[str]:
    """Detecta si una línea es un encabezado de sección."""
    clean = line.strip().lower()
    clean_no_punct = re.sub(r'[^\w\s]', '', clean)
    for section, keywords in _SECTION_KEYWORDS.items():
        for kw in keywords:
            if kw in clean_no_punct:
                return section
    return None


def _is_section_heading(line: str) -> bool:
    """Heurística para determinar si una línea es un encabezado."""
    stripped = line.strip()
    if not stripped or len(stripped) > 80:
        return False
    # Todo mayúsculas
    if stripped.isupper() and len(stripped) >= 3:
        return True
    # Detectada como sección conocida
    if _detect_section(stripped):
        return True
    # Termina en ":"
    if stripped.endswith(":") and len(stripped) <= 60:
        return True
    return False


def _is_contact_line(line: str) -> bool:
    """Detecta líneas con información de contacto."""
    patterns = [
        r'[\w.+-]+@[\w-]+\.\w+',            # email
        r'(?:\+?52|(?:\+?\d{1,3}))?[\s\-.]?\(?\d{3}\)?[\s\-.]?\d{3}[\s\-.]?\d{4}',  # teléfono
        r'linkedin\.com',
        r'github\.com',
        r'http[s]?://',
    ]
    for p in patterns:
        if re.search(p, line, re.IGNORECASE):
            return True
    return False


def _is_bullet(line: str) -> bool:
    return bool(re.match(r'^\s*[•\-*·▪‣◦►]\s+', line))


def _clean_bullet(line: str) -> str:
    return re.sub(r'^\s*[•\-*·▪‣◦►]\s+', '', line).strip()


def _is_date_range(line: str) -> bool:
    return bool(re.search(r'\b(19|20)\d{2}\b', line))


# ─── Parser de texto OCR ─────────────────────────────────────────────────────

class CVSection:
    def __init__(self, title: str):
        self.title = title
        self.items: list[str] = []  # bullet points o párrafos

    def add(self, text: str):
        if text.strip():
            self.items.append(text.strip())


def parse_cv_text(raw_text: str) -> dict:
    """
    Analiza el texto extraído por OCR y lo estructura en secciones.
    Devuelve un dict con: name, contact_lines, sections (list of CVSection).
    """
    lines = [l.rstrip() for l in raw_text.splitlines()]

    result = {
        "name": "",
        "contact_lines": [],
        "sections": [],
    }

    # Primer pase: detectar nombre (primeras líneas no vacías que no son contacto)
    header_done = False
    current_section: Optional[CVSection] = None
    current_text_lines: list[str] = []

    def flush_text():
        nonlocal current_text_lines
        if current_section and current_text_lines:
            paragraph = " ".join(current_text_lines).strip()
            if paragraph:
                current_section.add(paragraph)
        current_text_lines = []

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            flush_text()
            i += 1
            continue

        # Cabecera del CV (nombre y contacto)
        if not header_done:
            if not result["name"] and not _is_contact_line(stripped) and not _is_section_heading(stripped):
                result["name"] = stripped
                i += 1
                continue
            if _is_contact_line(stripped):
                result["contact_lines"].append(stripped)
                i += 1
                continue
            # Si llegamos a un encabezado de sección, terminamos la cabecera
            if _is_section_heading(stripped):
                header_done = True
                # no incrementamos i para que se procese abajo

        if not header_done:
            i += 1
            continue

        # Encabezado de sección
        if _is_section_heading(stripped):
            flush_text()
            label = stripped.rstrip(":").strip()
            current_section = CVSection(label)
            result["sections"].append(current_section)
            i += 1
            continue

        # Bullet point
        if _is_bullet(stripped):
            flush_text()
            if current_section is None:
                current_section = CVSection("Información")
                result["sections"].append(current_section)
            current_section.add(_clean_bullet(stripped))
            i += 1
            continue

        # Línea de texto normal
        if current_section is None:
            current_section = CVSection("Información")
            result["sections"].append(current_section)
        current_text_lines.append(stripped)
        i += 1

    flush_text()
    return result


# ─── Generador DOCX ───────────────────────────────────────────────────────────

def generate_ats_docx(raw_text: str) -> bytes:
    """
    Genera un archivo Word (.docx) optimizado para ATS a partir de texto OCR.
    Devuelve los bytes del archivo.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx no está instalado. Ejecuta: pip install python-docx"
        )

    parsed = parse_cv_text(raw_text)
    doc = Document()

    # ─── Márgenes estándar ATS (1 pulgada) ─────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ─── Estilo base ────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font  = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ─── Nombre ─────────────────────────────────────────────────────────────
    name = parsed.get("name", "").strip()
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # ─── Contacto ───────────────────────────────────────────────────────────
    contact = parsed.get("contact_lines", [])
    if contact:
        p = doc.add_paragraph(" | ".join(contact))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)

    if name or contact:
        doc.add_paragraph()  # espaciado

    # ─── Secciones ──────────────────────────────────────────────────────────
    for sec in parsed["sections"]:
        # Título de sección
        heading = doc.add_paragraph()
        run = heading.add_run(sec.title.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Línea separadora (usando caracteres — más compatibles con ATS que una línea de Word)
        sep = doc.add_paragraph("─" * 60)
        sep.runs[0].font.size = Pt(8)
        sep.paragraph_format.space_after = Pt(4)

        # Contenido
        for item in sec.items:
            p = doc.add_paragraph(style="List Bullet" if len(item) < 200 else "Normal")
            p.text = ""
            run = p.add_run(item)
            run.font.size = Pt(11)

        doc.add_paragraph()  # espacio entre secciones

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── Generador Texto Plano ATS ───────────────────────────────────────────────

def generate_ats_plain_text(raw_text: str) -> str:
    """
    Genera un CV en texto plano bien estructurado, ideal para copiar-pegar en formularios ATS.
    """
    parsed = parse_cv_text(raw_text)
    lines = []

    name = parsed.get("name", "").strip()
    if name:
        lines.append(name.upper())
        lines.append("=" * len(name))
        lines.append("")

    contact = parsed.get("contact_lines", [])
    for c in contact:
        lines.append(c)
    if contact:
        lines.append("")

    for sec in parsed["sections"]:
        lines.append(sec.title.upper())
        lines.append("-" * len(sec.title))
        for item in sec.items:
            # Decide si es un bullet o párrafo largo
            if len(item) < 200:
                lines.append(f"• {item}")
            else:
                lines.append(item)
        lines.append("")

    return "\n".join(lines)


# ─── Generador PDF ATS desde perfil estructurado ───────────────────────────────

def generate_ats_pdf_from_profile(profile: Dict[str, Any]) -> bytes:
    """
    Genera un PDF optimizado para ATS a partir de un perfil estructurado.
    """
    if not XHTML2PDF_AVAILABLE:
        raise RuntimeError(
            "xhtml2pdf no está instalado. Ejecuta: pip install xhtml2pdf"
        )

    # Build experience HTML
    experience_html = ''
    experience_data = profile.get("experience", [])
    if experience_data:
        experience_html = '<div class="section"><div class="section-title">Experiencia Profesional</div>'
        for job in experience_data:
            experience_html += f'''
                <div class="job">
                    <div class="job-title">{job.get("title", "")}</div>
                    <div class="job-company">{job.get("company", "")}</div>
                    {f'<div class="job-dates">{job.get("dates", "")}</div>' if job.get("dates") else ''}
                    {f'<ul>{"".join([f"<li>{desc}</li>" for desc in job.get("description", [])])}</ul>' if job.get("description") else ''}
                </div>
            '''
        experience_html += '</div>'
    elif profile.get("sections", {}).get("experience"):
        # Fallback to unstructured experience section
        experience_html = '<div class="section"><div class="section-title">Experiencia Profesional</div><ul>'
        experience_html += ''.join([f'<li>{exp}</li>' for exp in profile.get("sections", {}).get("experience", [])])
        experience_html += '</ul></div>'

    # Build skills HTML
    skills_html = ''
    skills = profile.get("skills", {})
    all_skills_flat = profile.get("all_skills_flat", [])
    if isinstance(skills, dict) and any(skills.values()):
        skills_html = '<div class="section"><div class="section-title">Habilidades</div><div class="skills-grid">'
        for cat, skill_list in skills.items():
            if skill_list:
                skills_html += f'<div><strong>{cat}:</strong> {", ".join(skill_list)}</div>'
        skills_html += '</div></div>'
    elif all_skills_flat:
        skills_html = f'<div class="section"><div class="section-title">Habilidades</div><p>{", ".join(all_skills_flat)}</p></div>'

    # Build languages HTML
    languages_html = ''
    languages = profile.get("languages_spoken", [])
    if languages:
        languages_html = f'<div class="section"><div class="section-title">Idiomas</div><ul>'
        languages_html += ''.join([f'<li>{lang}</li>' for lang in languages])
        languages_html += '</ul></div>'

    # Build HTML content
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>CV - {profile.get('name', 'Candidato')}</title>
        <style>
            @page {{
                size: A4 portrait;
                margin: 2.54cm;
            }}
            body {{
                font-family: Arial, Helvetica, sans-serif;
                font-size: 11pt;
                color: #1a1a2e;
                line-height: 1.5;
            }}
            .header {{
                text-align: center;
                margin-bottom: 20px;
            }}
            .name {{
                font-size: 18pt;
                font-weight: bold;
                margin-bottom: 10px;
            }}
            .contact {{
                font-size: 10pt;
            }}
            .section {{
                margin-bottom: 15px;
            }}
            .section-title {{
                font-size: 12pt;
                font-weight: bold;
                text-transform: uppercase;
                margin-bottom: 8px;
                border-bottom: 1px solid #1a1a2e;
                padding-bottom: 4px;
            }}
            .job {{
                margin-bottom: 12px;
            }}
            .job-title {{
                font-weight: bold;
            }}
            .job-company {{
                font-style: italic;
            }}
            .job-dates {{
                font-size: 10pt;
                color: #4a4a4a;
            }}
            ul {{
                margin: 5px 0 10px 20px;
                padding: 0;
            }}
            li {{
                margin-bottom: 5px;
            }}
            .skills-grid {{
                display: -pdf-box;
                -pdf-box-orient: horizontal;
            }}
            .skills-grid > div {{
                width: 50%;
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <div class="name">{profile.get('name', 'Candidato')}</div>
            <div class="contact">
                {profile.get('email', '')} | {profile.get('phone', '')} | {profile.get('location', '')}
                {f' | LinkedIn: {profile.get("linkedin", "")}' if profile.get("linkedin", "") else ''}
                {f' | GitHub: {profile.get("github", "")}' if profile.get("github", "") else ''}
            </div>
        </div>

        {f'<div class="section"><div class="section-title">Perfil Profesional</div><p>{profile.get("summary", "")}</p></div>' if profile.get("summary", "") else ''}

        {experience_html}

        {f'<div class="section"><div class="section-title">Educación</div><ul>{"".join([f"<li>{edu}</li>" for edu in profile.get("education", [])])}</ul></div>' if profile.get("education", []) else ''}

        {skills_html}

        {f'<div class="section"><div class="section-title">Certificaciones</div><ul>{"".join([f"<li>{cert}</li>" for cert in profile.get("certifications", [])])}</ul></div>' if profile.get("certifications", []) else ''}

        {languages_html}
    </body>
    </html>
    """

    buf = io.BytesIO()
    pisa.CreatePDF(
        io.BytesIO(html_content.encode('utf-8')),
        dest=buf,
        encoding='utf-8'
    )
    buf.seek(0)
    return buf.read()
